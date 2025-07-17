import os
import hashlib
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
import asyncio
import aiofiles

# File processing imports
try:
    import PyPDF2
    from PIL import Image
    import pytesseract
    from docx import Document
    import openpyxl
    import magic
except ImportError as e:
    print(f"Warning: Some file processing libraries not available: {e}")
    print("Install with: pip install PyPDF2 python-docx openpyxl pillow pytesseract python-magic")

from config import Config

@dataclass
class Attachment:
    id: str
    filename: str
    original_filename: str
    file_path: str
    file_type: str
    mime_type: str
    size: int
    upload_time: datetime
    extracted_text: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    thumbnail_path: Optional[str] = None

@dataclass
class AttachmentContext:
    attachment: Attachment
    relevance_score: float
    summary: str

class AttachmentManager:
    """Manages file attachments and content extraction"""
    
    def __init__(self, workspace_path: str = "."):
        self.workspace_path = Path(workspace_path).resolve()
        self.attachments_dir = self.workspace_path / ".mcp_attachments"
        self.thumbnails_dir = self.attachments_dir / "thumbnails"
        
        # Create directories
        self.attachments_dir.mkdir(exist_ok=True)
        self.thumbnails_dir.mkdir(exist_ok=True)
        
        # In-memory storage (in production, use a database)
        self.attachments: Dict[str, Attachment] = {}
        
        # Supported file types
        self.supported_types = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'document': ['.pdf', '.doc', '.docx', '.txt', '.md', '.rtf'],
            'spreadsheet': ['.xls', '.xlsx', '.csv'],
            'code': ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.h', '.css', '.html'],
            'archive': ['.zip', '.tar', '.gz', '.rar'],
            'other': []
        }
    
    async def upload_attachment(
        self, 
        file_content: bytes, 
        filename: str,
        extract_content: bool = True
    ) -> Attachment:
        """Upload and process an attachment"""
        
        # Generate unique ID
        file_hash = hashlib.md5(file_content).hexdigest()
        attachment_id = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file_hash[:8]}"
        
        # Determine file type and mime type
        file_extension = Path(filename).suffix.lower()
        mime_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
        file_type = self._get_file_type(file_extension)
        
        # Save file
        safe_filename = self._sanitize_filename(filename)
        file_path = self.attachments_dir / f"{attachment_id}_{safe_filename}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Create attachment object
        attachment = Attachment(
            id=attachment_id,
            filename=safe_filename,
            original_filename=filename,
            file_path=str(file_path),
            file_type=file_type,
            mime_type=mime_type,
            size=len(file_content),
            upload_time=datetime.now()
        )
        
        # Extract content if requested
        if extract_content:
            try:
                extracted_text, metadata = await self._extract_content(attachment)
                attachment.extracted_text = extracted_text
                attachment.metadata = metadata
            except Exception as e:
                print(f"Failed to extract content from {filename}: {e}")
        
        # Generate thumbnail for images
        if file_type == 'image':
            try:
                thumbnail_path = await self._generate_thumbnail(attachment)
                attachment.thumbnail_path = thumbnail_path
            except Exception as e:
                print(f"Failed to generate thumbnail for {filename}: {e}")
        
        # Store attachment
        self.attachments[attachment_id] = attachment
        
        return attachment
    
    async def _extract_content(self, attachment: Attachment) -> Tuple[Optional[str], Optional[Dict]]:
        """Extract text content from various file types"""
        
        file_path = Path(attachment.file_path)
        file_type = attachment.file_type
        
        try:
            if file_type == 'document':
                return await self._extract_document_content(file_path)
            elif file_type == 'image':
                return await self._extract_image_content(file_path)
            elif file_type == 'spreadsheet':
                return await self._extract_spreadsheet_content(file_path)
            elif file_type == 'code':
                return await self._extract_text_content(file_path)
            else:
                # Try to extract as text
                return await self._extract_text_content(file_path)
        except Exception as e:
            print(f"Error extracting content: {e}")
            return None, None
    
    async def _extract_document_content(self, file_path: Path) -> Tuple[Optional[str], Optional[Dict]]:
        """Extract content from document files"""
        
        if file_path.suffix.lower() == '.pdf':
            return await self._extract_pdf_content(file_path)
        elif file_path.suffix.lower() in ['.doc', '.docx']:
            return await self._extract_docx_content(file_path)
        elif file_path.suffix.lower() in ['.txt', '.md']:
            return await self._extract_text_content(file_path)
        
        return None, None
    
    async def _extract_pdf_content(self, file_path: Path) -> Tuple[Optional[str], Optional[Dict]]:
        """Extract text from PDF files"""
        try:
            text_content = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        print(f"Error extracting page {page_num + 1}: {e}")
            
            return '\n\n'.join(text_content), metadata
            
        except Exception as e:
            print(f"Error extracting PDF content: {e}")
            return None, None
    
    async def _extract_docx_content(self, file_path: Path) -> Tuple[Optional[str], Optional[Dict]]:
        """Extract text from Word documents"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract metadata
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or ''
            }
            
            return '\n\n'.join(text_content), metadata
            
        except Exception as e:
            print(f"Error extracting DOCX content: {e}")
            return None, None
    
    async def _extract_image_content(self, file_path: Path) -> Tuple[Optional[str], Optional[Dict]]:
        """Extract text from images using OCR"""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Extract text using OCR
            extracted_text = pytesseract.image_to_string(image)
            
            # Get image metadata
            metadata = {
                'width': image.width,
                'height': image.height,
                'format': image.format,
                'mode': image.mode
            }
            
            # Add EXIF data if available
            if hasattr(image, '_getexif') and image._getexif():
                metadata['exif'] = dict(image._getexif())
            
            return extracted_text.strip() if extracted_text.strip() else None, metadata
            
        except Exception as e:
            print(f"Error extracting image content: {e}")
            return None, None
    
    async def _extract_spreadsheet_content(self, file_path: Path) -> Tuple[Optional[str], Optional[Dict]]:
        """Extract content from spreadsheet files"""
        try:
            if file_path.suffix.lower() in ['.xls', '.xlsx']:
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                
                text_content = []
                metadata = {'sheets': list(workbook.sheetnames)}
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_content = [f"--- Sheet: {sheet_name} ---"]
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                        if row_text.strip():
                            sheet_content.append(row_text)
                    
                    text_content.extend(sheet_content)
                
                return '\n'.join(text_content), metadata
                
        except Exception as e:
            print(f"Error extracting spreadsheet content: {e}")
            return None, None
    
    async def _extract_text_content(self, file_path: Path) -> Tuple[Optional[str], Optional[Dict]]:
        """Extract content from text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
            
            metadata = {
                'lines': len(content.splitlines()),
                'characters': len(content),
                'encoding': 'utf-8'
            }
            
            return content, metadata
            
        except Exception as e:
            print(f"Error extracting text content: {e}")
            return None, None
    
    async def _generate_thumbnail(self, attachment: Attachment) -> Optional[str]:
        """Generate thumbnail for image files"""
        try:
            image = Image.open(attachment.file_path)
            
            # Create thumbnail
            thumbnail_size = (200, 200)
            image.thumbnail(thumbnail_size, Image.Resampling.LANCZOS)
            
            # Save thumbnail
            thumbnail_filename = f"{attachment.id}_thumb.jpg"
            thumbnail_path = self.thumbnails_dir / thumbnail_filename
            
            # Convert to RGB if necessary
            if image.mode in ('RGBA', 'LA', 'P'):
                image = image.convert('RGB')
            
            image.save(thumbnail_path, 'JPEG', quality=85)
            
            return str(thumbnail_path)
            
        except Exception as e:
            print(f"Error generating thumbnail: {e}")
            return None
    
    def _get_file_type(self, extension: str) -> str:
        """Determine file type category from extension"""
        for file_type, extensions in self.supported_types.items():
            if extension in extensions:
                return file_type
        return 'other'
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe storage"""
        # Remove or replace dangerous characters
        safe_chars = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789.-_"
        sanitized = ''.join(c if c in safe_chars else '_' for c in filename)
        
        # Limit length
        if len(sanitized) > 100:
            name, ext = os.path.splitext(sanitized)
            sanitized = name[:95] + ext
        
        return sanitized
    
    def get_attachment(self, attachment_id: str) -> Optional[Attachment]:
        """Get attachment by ID"""
        return self.attachments.get(attachment_id)
    
    def list_attachments(self) -> List[Attachment]:
        """List all attachments"""
        return list(self.attachments.values())
    
    async def delete_attachment(self, attachment_id: str) -> bool:
        """Delete an attachment"""
        attachment = self.attachments.get(attachment_id)
        if not attachment:
            return False
        
        try:
            # Delete main file
            if os.path.exists(attachment.file_path):
                os.remove(attachment.file_path)
            
            # Delete thumbnail if exists
            if attachment.thumbnail_path and os.path.exists(attachment.thumbnail_path):
                os.remove(attachment.thumbnail_path)
            
            # Remove from memory
            del self.attachments[attachment_id]
            
            return True
            
        except Exception as e:
            print(f"Error deleting attachment: {e}")
            return False
    
    def get_relevant_attachments(self, query: str, limit: int = 5) -> List[AttachmentContext]:
        """Get attachments relevant to a query"""
        relevant = []
        
        query_lower = query.lower()
        
        for attachment in self.attachments.values():
            relevance_score = 0.0
            
            # Check filename relevance
            if query_lower in attachment.filename.lower():
                relevance_score += 0.3
            
            # Check extracted text relevance
            if attachment.extracted_text:
                text_lower = attachment.extracted_text.lower()
                if query_lower in text_lower:
                    relevance_score += 0.7
                    # Count occurrences for better scoring
                    occurrences = text_lower.count(query_lower)
                    relevance_score += min(occurrences * 0.1, 0.5)
            
            if relevance_score > 0:
                # Create summary
                summary = self._create_attachment_summary(attachment, query)
                
                relevant.append(AttachmentContext(
                    attachment=attachment,
                    relevance_score=relevance_score,
                    summary=summary
                ))
        
        # Sort by relevance and return top results
        relevant.sort(key=lambda x: x.relevance_score, reverse=True)
        return relevant[:limit]
    
    def _create_attachment_summary(self, attachment: Attachment, query: str) -> str:
        """Create a summary of attachment content relevant to query"""
        summary_parts = [f"File: {attachment.original_filename}"]
        
        if attachment.metadata:
            if 'title' in attachment.metadata and attachment.metadata['title']:
                summary_parts.append(f"Title: {attachment.metadata['title']}")
        
        if attachment.extracted_text:
            # Find relevant excerpts
            text = attachment.extracted_text
            query_lower = query.lower()
            
            # Simple excerpt extraction
            sentences = text.split('.')
            relevant_sentences = [s.strip() for s in sentences if query_lower in s.lower()]
            
            if relevant_sentences:
                excerpt = '. '.join(relevant_sentences[:3])
                if len(excerpt) > 200:
                    excerpt = excerpt[:200] + "..."
                summary_parts.append(f"Relevant content: {excerpt}")
        
        return '\n'.join(summary_parts)
